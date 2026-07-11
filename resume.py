from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc=Document()
style=doc.styles['Normal']
style.font.name='Calibri'
style.font.size=Pt(10)

t=doc.add_heading('Ram Baranwal',1)
t.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
p=doc.add_paragraph()
p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run("Azamgarh, Uttar Pradesh 276125\n")
p.add_run("+91 9919799661 | rambaranwal644@gmail.com | linkedin.com/in/ram-baranwal | github.com/RamBaranwal")

doc.add_heading('Professional Summary', level=2)
doc.add_paragraph(
"Computer Science undergraduate with hands-on experience in Java, MERN Stack, AWS, Docker, CI/CD, and cloud-native application deployment. "
"Strong foundation in Core Java, Object-Oriented Programming, JDBC, SQL, REST APIs, Data Structures & Algorithms, and Linux. "
"Experienced in building scalable web applications and automated deployment pipelines using GitHub Actions and AWS EC2. "
"Seeking Software Engineer or Java Developer opportunities."
)

doc.add_heading('Projects',2)
projects=[
("Cloud-Native Game Deployment Pipeline using DevOps | Node.js, React.js, Docker, GitHub Actions, AWS | Mar 2026",[
"Built a full-stack Snake & Fruit Game using Vanilla JavaScript and React.js, then containerized both applications with Docker for consistent deployments.",
"Engineered CI/CD pipelines using GitHub Actions to automatically build, test, and publish Docker images, reducing manual deployment effort by approximately 30 minutes per release.",
"Deployed on AWS EC2 with Docker, Nginx reverse proxy, and SSH automation to ensure reliable public access and zero-touch deployment.",
"GitHub: https://github.com/RamBaranwal/fruitAndSnake"]),
("Hospital Management System | MongoDB, Express.js, React.js, REST APIs | Feb 2026",[
"Developed secure role-based authentication for Doctors and Patients with MongoDB integration.",
"Designed RESTful APIs for appointment scheduling, doctor management, and patient records.",
"Implemented complete appointment workflow with real-time status synchronization.",
"GitHub: https://github.com/RamBaranwal/hospital"]),
("NexDo Deployment on AWS with Docker | MERN, AWS EC2, Docker | Jan 2026",[
"Containerized a MERN To-Do application using Docker for consistent development and production environments.",
"Configured and deployed the application on AWS EC2 Ubuntu with optimized Docker networking and security.",
"Established Git-based deployment workflows to simplify cloud operations.",
"GitHub: https://github.com/RamBaranwal/aws-first-project"]),
("LibSync – Digital Library Solution | Java, JDBC, MySQL, Apache NetBeans | Jul 2025",[
"Built a Java Swing desktop application for student, book, issue, and return management.",
"Integrated JDBC with MySQL to perform secure CRUD operations and transaction management.",
"Applied OOP principles, exception handling, authentication, and input validation.",
"GitHub: https://github.com/RamBaranwal/Updated-LMS"])
]
for title, bullets in projects:
    doc.add_heading(title,3)
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Technical Skills',2)
doc.add_paragraph(
"Languages: Java, JavaScript, Python\n"
"Core Java: OOP, Collections Framework, Exception Handling, JDBC, Multithreading (Basics)\n"
"Web & Frameworks: React.js, Node.js, Express.js, REST APIs\n"
"Databases: MySQL, MongoDB\n"
"Cloud & DevOps: AWS EC2, Docker, GitHub Actions, CI/CD, Nginx, Linux, Ubuntu\n"
"Tools: Git, GitHub, Docker Hub, Apache NetBeans, VS Code, Postman\n"
"Core Concepts: Data Structures & Algorithms, DBMS, Operating Systems, Computer Networks, Software Engineering, Cloud Computing"
)

doc.add_heading('Relevant Coursework',2)
doc.add_paragraph("Data Structures & Algorithms • Object-Oriented Programming • Database Management Systems • Operating Systems • Computer Networks • Software Engineering • Cloud Computing")

doc.add_heading('Achievements',2)
for a in [
"Solved 500+ Data Structures & Algorithms problems.",
"Earned HackerRank Java Gold Badge.",
"Participated in DUHack 3.0 Hackathon.",
"Completed AWS Academy Cloud Architecting and Cloud Foundations training."
]:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('Certificates',2)
doc.add_paragraph(
"AWS Academy Graduate – Cloud Architecting Training Badge (Jan 2026)\n"
"Oracle Cloud Infrastructure 2025 Certified AI Foundations Associate (Oct 2025)\n"
"AWS Academy Graduate – Cloud Foundations Training Badge (Jun 2025)"
)

doc.add_heading('Education',2)
doc.add_paragraph(
"Lovely Professional University, Phagwara, Punjab\n"
"B.Tech in Computer Science & Engineering | CGPA: 7.77 | Aug 2023 – Present\n\n"
"Decent School, Kota, Rajasthan\n"
"Class XII – 72.60% | 2023\n\n"
"Sunbeam School, Azamgarh, Uttar Pradesh\n"
"Class X – 86.40% | 2020"
)

out="/mnt/data/Ram_Baranwal_Updated_ATS_Resume.docx"
doc.save(out)
print(out)
